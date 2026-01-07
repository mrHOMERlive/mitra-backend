from io import BytesIO
from typing import Dict
from uuid import UUID
from docx import Document
import pymorphy3
from app.models import NDAType, FieldsENG, FieldsRuEn
from app.services.minio_service import minio_service


class DOCXGenerator:
    TEMPLATE_MAP = {
        NDAType.ENG: "PT MITRA - NDA_eng.docx",
        NDAType.RU_EN: "PT MITRA - NDA_rus_eng.docx"
    }

    FIELD_MAPPING_ENG = {
        "POINT 1": "effective_date",
        "POINT 2": "company_name",
        "POINT 3": "country",
        "POINT 4": "registration_number",
        "POINT 5": "signatory_name",
        "POINT 5.1": "signatory_title",
        "POINT 6": "address",
        "POINT 7": "email"
    }

    FIELD_MAPPING_RU_EN = {
        "POINT 1": "effective_date",
        "POINT 2": "company_name_en",
        "POINT 3": "company_name_ru",
        "POINT 4": "country_en",
        "POINT 5": "country_ru",
        "POINT 6": "registration_number",
        "POINT 7": "signatory_name_en",
        "POINT 7.1": "signatory_title_en",
        "POINT 8": "signatory_name_ru",
        "POINT 9": "address_en",
        "POINT 10": "address_ru",
        "POINT 11": "email"
    }

    def __init__(self):
        self.morph = pymorphy3.MorphAnalyzer()

    def _get_field_mapping(self, nda_type: NDAType) -> Dict[str, str]:
        if nda_type == NDAType.ENG:
            return self.FIELD_MAPPING_ENG
        elif nda_type == NDAType.RU_EN:
            return self.FIELD_MAPPING_RU_EN
        else:
            raise ValueError(f"Unknown NDA type: {nda_type}")

    def _to_genitive(self, text: str) -> str:
        if not text:
            return text
        
        words = text.split()
        inflected_words = []
        
        for word in words:
            parses = self.morph.parse(word)
            target_parse = None
            
            for p in parses:
                if 'nomn' in p.tag:
                    target_parse = p
                    break
            
            if not target_parse:
                target_parse = parses[0]
            
            inflected = target_parse.inflect({'gent'})
            
            if inflected:
                result_word = inflected.word
                if word.istitle():
                    result_word = result_word.capitalize()
                elif word.isupper():
                    result_word = result_word.upper()
                inflected_words.append(result_word)
            else:
                inflected_words.append(word)
                
        return " ".join(inflected_words)

    def _replace_placeholders(self, doc: Document, fields: Dict, mapping: Dict[str, str]) -> None:
        replacements = {
            f"[{placeholder}]": str(fields[field_name])
            for placeholder, field_name in mapping.items()
            if field_name in fields and fields[field_name] is not None
        }

        def replace_in_paragraph(paragraph):
            if not paragraph.runs:
                return
            
            full_text = ''.join(run.text for run in paragraph.runs)
            
            modified_text = full_text
            for placeholder, value in replacements.items():
                modified_text = modified_text.replace(placeholder, value)
            
            if modified_text == full_text:
                return
            
            char_to_run = []
            for run_idx, run in enumerate(paragraph.runs):
                char_to_run.extend([run_idx] * len(run.text))
            
            position_map = []
            old_pos = 0
            new_pos = 0
            
            sorted_placeholders = sorted(replacements.keys(), key=len, reverse=True)
            
            while old_pos < len(full_text):
                placeholder_found = None
                for placeholder in sorted_placeholders:
                    if full_text[old_pos:old_pos + len(placeholder)] == placeholder:
                        placeholder_found = placeholder
                        break
                
                if placeholder_found:
                    replacement = replacements[placeholder_found]
                    run_idx = char_to_run[old_pos] if old_pos < len(char_to_run) else 0
                    
                    for _ in range(len(replacement)):
                        position_map.append(run_idx)
                        new_pos += 1
                    
                    old_pos += len(placeholder_found)
                else:
                    run_idx = char_to_run[old_pos] if old_pos < len(char_to_run) else 0
                    position_map.append(run_idx)
                    old_pos += 1
                    new_pos += 1
            
            new_run_texts = [''] * len(paragraph.runs)
            for i, char in enumerate(modified_text):
                if i < len(position_map):
                    run_idx = position_map[i]
                    new_run_texts[run_idx] += char
                else:
                    new_run_texts[0] += char
            
            for i, run in enumerate(paragraph.runs):
                run.text = new_run_texts[i]

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

    def generate(self, nda_id: UUID, nda_type: NDAType, fields: Dict) -> bytes:
        template_name = self.TEMPLATE_MAP.get(nda_type)
        if not template_name:
            raise ValueError(f"No template found for NDA type: {nda_type}")

        template_bytes = minio_service.get_template(template_name)
        
        doc = Document(BytesIO(template_bytes))
        
        field_mapping = self._get_field_mapping(nda_type)
        
        processed_fields = fields.copy()
        
        if nda_type == NDAType.RU_EN and "signatory_name_ru" in processed_fields:
            processed_fields["signatory_name_ru"] = self._to_genitive(processed_fields["signatory_name_ru"])
        
        self._replace_placeholders(doc, processed_fields, field_mapping)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.read()


docx_generator = DOCXGenerator()